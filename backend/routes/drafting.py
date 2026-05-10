"""backend/routes/drafting.py — All 20 document types"""

import io
from flask import Blueprint, render_template, request, jsonify, send_file, current_app
from flask_login import login_required, current_user
from backend.database import get_db
from backend.ai_service import generate_document

drafting_bp = Blueprint("drafting", __name__)

# ── All 20 document types from the project spec ──────────────────────────────
DOC_TYPES = [
    "Legal Notice",
    "Affidavit",
    "Rental / Lease Agreement",
    "Employment Contract",
    "Non-Disclosure Agreement (NDA)",
    "Sale Deed",
    "Power of Attorney",
    "Partnership Deed",
    "Memorandum of Understanding (MoU)",
    "Demand Notice",
    "Complaint / Petition",
    "Bail Application",
    "Writ Petition",
    "Divorce Petition",
    "RTI Application",
    "Settlement Agreement",
    "Service Agreement",
    "Vendor / Supplier Agreement",
    "Will / Testament",
    "Loan Agreement",
]

# ── Per-type field definitions ────────────────────────────────────────────────
DOC_FIELDS = {
    "Legal Notice": [
        ("sender",      "Sender Full Name"),
        ("recipient",   "Recipient Full Name / Company"),
        ("address",     "Recipient Address"),
        ("subject",     "Subject of Notice"),
        ("grievance",   "Details of Grievance / Demand"),
        ("relief",      "Relief Sought"),
        ("deadline",    "Response Deadline (days)"),
        ("lawyer",      "Advocate Name (if any)"),
    ],
    "Affidavit": [
        ("deponent",    "Deponent Full Name"),
        ("age",         "Age"),
        ("address",     "Address"),
        ("id_proof",    "ID Proof Type (Aadhaar/Passport/etc.)"),
        ("statement",   "Statement / Purpose of Affidavit"),
        ("location",    "Place of Execution"),
        ("date",        "Date"),
    ],
    "Rental / Lease Agreement": [
        ("landlord",    "Landlord Full Name"),
        ("tenant",      "Tenant Full Name"),
        ("property",    "Property Full Address"),
        ("rent",        "Monthly Rent (₹)"),
        ("deposit",     "Security Deposit (₹)"),
        ("start_date",  "Lease Start Date"),
        ("duration",    "Lease Duration (months)"),
        ("purpose",     "Purpose (Residential/Commercial)"),
    ],
    "Employment Contract": [
        ("employer",    "Employer / Company Name"),
        ("employee",    "Employee Full Name"),
        ("designation", "Job Title / Designation"),
        ("department",  "Department"),
        ("start_date",  "Start Date"),
        ("salary",      "Monthly Gross Salary (₹)"),
        ("work_hours",  "Weekly Work Hours"),
        ("notice_period","Notice Period"),
        ("location",    "Work Location"),
        ("probation",   "Probation Period"),
    ],
    "Non-Disclosure Agreement (NDA)": [
        ("party_a",     "Disclosing Party Name"),
        ("party_b",     "Receiving Party Name"),
        ("purpose",     "Purpose of Disclosure"),
        ("duration",    "Confidentiality Duration"),
        ("jurisdiction","Jurisdiction / State"),
        ("date",        "Effective Date"),
    ],
    "Sale Deed": [
        ("seller",      "Seller Full Name"),
        ("buyer",       "Buyer Full Name"),
        ("property",    "Property Description / Address"),
        ("area",        "Area (sq. ft. / sq. m.)"),
        ("price",       "Sale Consideration (₹)"),
        ("advance",     "Advance Paid (₹)"),
        ("date",        "Date of Sale"),
        ("registrar",   "Sub-Registrar Office"),
    ],
    "Power of Attorney": [
        ("principal",   "Principal (Grantor) Full Name"),
        ("agent",       "Agent (Attorney) Full Name"),
        ("purpose",     "Purpose / Powers Granted"),
        ("scope",       "Scope of Authority"),
        ("duration",    "Duration / Expiry Date"),
        ("location",    "Place of Execution"),
    ],
    "Partnership Deed": [
        ("firm_name",   "Partnership Firm Name"),
        ("partner_a",   "Partner A Full Name"),
        ("partner_b",   "Partner B Full Name"),
        ("partner_c",   "Partner C Full Name (optional)"),
        ("business",    "Nature of Business"),
        ("capital",     "Total Capital (₹)"),
        ("profit_split","Profit/Loss Ratio"),
        ("start_date",  "Commencement Date"),
        ("address",     "Principal Place of Business"),
    ],
    "Memorandum of Understanding (MoU)": [
        ("party_a",     "Party A Name / Organisation"),
        ("party_b",     "Party B Name / Organisation"),
        ("purpose",     "Purpose of MoU"),
        ("scope",       "Scope of Collaboration"),
        ("duration",    "Duration"),
        ("date",        "Effective Date"),
        ("location",    "Place of Signing"),
    ],
    "Demand Notice": [
        ("sender",      "Sender / Creditor Name"),
        ("recipient",   "Recipient / Debtor Name"),
        ("amount",      "Amount Demanded (₹)"),
        ("reason",      "Reason for Demand"),
        ("due_date",    "Original Due Date"),
        ("deadline",    "Payment Deadline (days)"),
        ("account",     "Bank Account / Payment Details"),
    ],
    "Complaint / Petition": [
        ("petitioner",  "Petitioner Full Name"),
        ("respondent",  "Respondent / Opposite Party"),
        ("court",       "Court / Authority Name"),
        ("subject",     "Subject of Complaint"),
        ("facts",       "Brief Statement of Facts"),
        ("relief",      "Relief Prayed For"),
        ("date",        "Date of Filing"),
    ],
    "Bail Application": [
        ("applicant",   "Applicant Full Name"),
        ("case_no",     "FIR / Case Number"),
        ("police_station","Police Station"),
        ("offence",     "Offence / Section"),
        ("date_arrest", "Date of Arrest"),
        ("court",       "Court Name"),
        ("grounds",     "Grounds for Bail"),
        ("surety",      "Proposed Surety Details"),
    ],
    "Writ Petition": [
        ("petitioner",  "Petitioner Full Name"),
        ("respondent",  "Respondent (Govt. Body / Authority)"),
        ("court",       "High Court / Supreme Court"),
        ("writ_type",   "Type of Writ (Habeas Corpus / Mandamus / etc.)"),
        ("facts",       "Brief Facts"),
        ("grounds",     "Grounds for the Writ"),
        ("relief",      "Relief Sought"),
    ],
    "Divorce Petition": [
        ("petitioner",  "Petitioner Full Name"),
        ("respondent",  "Respondent (Spouse) Full Name"),
        ("marriage_date","Date of Marriage"),
        ("marriage_place","Place of Marriage"),
        ("grounds",     "Grounds for Divorce"),
        ("children",    "Children Details (if any)"),
        ("court",       "Family Court"),
        ("relief",      "Relief / Custody Sought"),
    ],
    "RTI Application": [
        ("applicant",   "Applicant Full Name"),
        ("address",     "Applicant Address"),
        ("authority",   "Public Authority / Department"),
        ("information", "Information Sought (detailed)"),
        ("period",      "Period of Information"),
        ("format",      "Preferred Format (certified copy / inspection)"),
        ("date",        "Date of Application"),
    ],
    "Settlement Agreement": [
        ("party_a",     "Party A Full Name"),
        ("party_b",     "Party B Full Name"),
        ("dispute",     "Nature of Dispute"),
        ("terms",       "Settlement Terms"),
        ("amount",      "Settlement Amount (₹) if any"),
        ("date",        "Date of Settlement"),
        ("jurisdiction","Jurisdiction"),
    ],
    "Service Agreement": [
        ("client",      "Client Name / Company"),
        ("provider",    "Service Provider Name / Company"),
        ("services",    "Services to Be Provided"),
        ("fee",         "Service Fee (₹)"),
        ("duration",    "Agreement Duration"),
        ("payment_terms","Payment Terms"),
        ("start_date",  "Start Date"),
        ("location",    "Place of Service"),
    ],
    "Vendor / Supplier Agreement": [
        ("buyer",       "Buyer / Company Name"),
        ("vendor",      "Vendor / Supplier Name"),
        ("goods",       "Goods / Services Supplied"),
        ("price",       "Price / Rate (₹)"),
        ("payment_terms","Payment Terms"),
        ("delivery",    "Delivery Terms"),
        ("duration",    "Agreement Duration"),
        ("warranty",    "Warranty / Guarantee Terms"),
    ],
    "Will / Testament": [
        ("testator",    "Testator Full Name"),
        ("age",         "Age"),
        ("address",     "Address"),
        ("id_proof",    "ID Proof Type"),
        ("assets",      "Assets / Properties to Bequeath"),
        ("beneficiaries","Beneficiaries and Their Share"),
        ("executor",    "Executor Full Name"),
        ("date",        "Date of Will"),
        ("witness_a",   "Witness 1 Name"),
        ("witness_b",   "Witness 2 Name"),
    ],
    "Loan Agreement": [
        ("lender",      "Lender Full Name / Company"),
        ("borrower",    "Borrower Full Name"),
        ("amount",      "Loan Amount (₹)"),
        ("interest",    "Interest Rate (% per annum)"),
        ("tenure",      "Loan Tenure"),
        ("repayment",   "Repayment Schedule"),
        ("purpose",     "Purpose of Loan"),
        ("security",    "Security / Collateral (if any)"),
        ("date",        "Date of Agreement"),
    ],
}


@drafting_bp.route("/")
@login_required
def index():
    db   = get_db(current_app.config["DATABASE"])
    docs = db.execute(
        "SELECT * FROM documents WHERE user_id=? AND is_archived=0 ORDER BY created_at DESC",
        (current_user.id,)
    ).fetchall()
    db.close()
    return render_template("drafting.html", doc_types=DOC_TYPES, docs=docs, doc_fields=DOC_FIELDS)


@drafting_bp.route("/fields/<path:doc_type>")
@login_required
def get_fields(doc_type):
    fields = DOC_FIELDS.get(doc_type, [
        ("party_a", "Party A Name"), ("party_b", "Party B Name"),
        ("purpose", "Purpose"),      ("date",    "Effective Date"),
    ])
    return jsonify({"fields": [{"key": k, "label": l} for k, l in fields]})


@drafting_bp.route("/generate", methods=["POST"])
@login_required
def generate():
    data     = request.get_json()
    doc_type = data.get("doc_type", "").strip()
    fields   = data.get("fields", {})
    if not doc_type:
        return jsonify({"error": "Document type required"}), 400

    content = generate_document(doc_type, fields, current_app.get_api_key())

    # Build title from primary field
    primary = (fields.get("employer") or fields.get("sender") or fields.get("petitioner")
               or fields.get("testator") or fields.get("party_a") or fields.get("landlord")
               or fields.get("lender") or "Document")
    title = f"{doc_type} — {primary}"[:120]

    db  = get_db(current_app.config["DATABASE"])
    cur = db.execute(
        "INSERT INTO documents (user_id,title,doc_type,content) VALUES (?,?,?,?)",
        (current_user.id, title, doc_type, content)
    )
    db.commit()
    doc_id = cur.lastrowid
    db.close()

    return jsonify({"content": content, "doc_id": doc_id, "title": title})


@drafting_bp.route("/download/<int:doc_id>")
@login_required
def download(doc_id):
    db  = get_db(current_app.config["DATABASE"])
    doc = db.execute(
        "SELECT * FROM documents WHERE id=? AND user_id=?",
        (doc_id, current_user.id)
    ).fetchone()
    db.close()
    if not doc:
        return "Not found", 404
    buf = _make_docx(doc["title"], doc["content"])
    return send_file(buf, as_attachment=True,
                     download_name=f"{doc['doc_type'].replace('/','_')}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@drafting_bp.route("/pin/<int:doc_id>", methods=["POST"])
@login_required
def pin(doc_id):
    db = get_db(current_app.config["DATABASE"])
    row = db.execute("SELECT is_pinned FROM documents WHERE id=? AND user_id=?",
                     (doc_id, current_user.id)).fetchone()
    if not row:
        db.close(); return jsonify({"error": "Not found"}), 404
    new_val = 0 if row["is_pinned"] else 1
    db.execute("UPDATE documents SET is_pinned=? WHERE id=?", (new_val, doc_id))
    db.commit(); db.close()
    return jsonify({"ok": True, "pinned": bool(new_val)})


@drafting_bp.route("/archive/<int:doc_id>", methods=["POST"])
@login_required
def archive(doc_id):
    db = get_db(current_app.config["DATABASE"])
    row = db.execute("SELECT is_archived FROM documents WHERE id=? AND user_id=?",
                     (doc_id, current_user.id)).fetchone()
    if not row:
        db.close(); return jsonify({"error": "Not found"}), 404
    new_val = 0 if row["is_archived"] else 1
    db.execute("UPDATE documents SET is_archived=? WHERE id=?", (new_val, doc_id))
    db.commit(); db.close()
    return jsonify({"ok": True, "archived": bool(new_val)})


@drafting_bp.route("/rename/<int:doc_id>", methods=["POST"])
@login_required
def rename(doc_id):
    data  = request.get_json()
    title = data.get("title", "").strip()
    if not title:
        return jsonify({"error": "Title required"}), 400
    db = get_db(current_app.config["DATABASE"])
    db.execute("UPDATE documents SET title=? WHERE id=? AND user_id=?",
               (title, doc_id, current_user.id))
    db.commit(); db.close()
    return jsonify({"ok": True})


@drafting_bp.route("/delete/<int:doc_id>", methods=["POST"])
@login_required
def delete(doc_id):
    db = get_db(current_app.config["DATABASE"])
    db.execute("DELETE FROM documents WHERE id=? AND user_id=?", (doc_id, current_user.id))
    db.commit(); db.close()
    return jsonify({"ok": True})


def _make_docx(title: str, content: str) -> __import__("io").BytesIO:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()
    h   = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for line in content.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("## ") or (line.isupper() and len(line) > 4):
            doc.add_heading(line.lstrip("# "), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.lstrip("# "), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@drafting_bp.route("/view/<int:doc_id>")
@login_required
def view_doc(doc_id):
    """Return document content as JSON so it can be loaded into the preview."""
    db  = get_db(current_app.config["DATABASE"])
    doc = db.execute("SELECT * FROM documents WHERE id=? AND user_id=?",
                     (doc_id, current_user.id)).fetchone()
    db.close()
    if not doc:
        return jsonify({"error": "Not found"}), 404
    return jsonify({"title": doc["title"], "doc_type": doc["doc_type"],
                    "content": doc["content"], "doc_id": doc["id"]})


@drafting_bp.route("/update/<int:doc_id>", methods=["POST"])
@login_required
def update_doc(doc_id):
    """Save edited document content."""
    data    = request.get_json()
    content = data.get("content", "").strip()
    db      = get_db(current_app.config["DATABASE"])
    db.execute("UPDATE documents SET content=?, updated_at=datetime('now') WHERE id=? AND user_id=?",
               (content, doc_id, current_user.id))
    db.commit(); db.close()
    return jsonify({"ok": True})


@drafting_bp.route("/download-pdf/<int:doc_id>")
@login_required
def download_pdf(doc_id):
    """Download document as PDF."""
    db  = get_db(current_app.config["DATABASE"])
    doc = db.execute("SELECT * FROM documents WHERE id=? AND user_id=?",
                     (doc_id, current_user.id)).fetchone()
    db.close()
    if not doc:
        return "Not found", 404
    buf = _make_pdf(doc["title"], doc["content"])
    return send_file(buf, as_attachment=True,
                     download_name=f"{doc['doc_type'].replace('/','_')}.pdf",
                     mimetype="application/pdf")


def _make_pdf(title: str, content: str):
    """Generate a PDF using fpdf2."""
    import io
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.set_margins(20, 20, 20)
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_xy(20, 20)
        pdf.multi_cell(0, 8, title, align="C")
        pdf.ln(6)
        pdf.set_font("Helvetica", "", 10)
        for line in content.split("\n"):
            line = line.rstrip()
            if not line:
                pdf.ln(4)
            elif line.isupper() and len(line) > 3:
                pdf.set_font("Helvetica", "B", 11)
                pdf.multi_cell(0, 7, line)
                pdf.set_font("Helvetica", "", 10)
            else:
                pdf.multi_cell(0, 6, line)
        out = pdf.output()
        return io.BytesIO(out)
    except Exception as e:
        # Fallback: plain text wrapped in minimal PDF marker
        import io
        buf = io.BytesIO()
        buf.write(f"{title}\n\n{content}".encode("utf-8", errors="replace"))
        buf.seek(0)
        return buf
